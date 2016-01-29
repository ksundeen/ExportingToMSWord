
# Program Purpose:
"""This script uses the python-docx module to create, access, and update
existing Microsoft Word documents for documenting metadata of hundreds of
Esri feature classes contained within an enterprise SDE geodatabase environment.

Specifically, It sets your development space for editing in a 64-bit OS and sets
your python variables (since the python-docx module is 64-bit). Using the docx
module, you access your SDE GIS database to grab attributes within each feature
class within each feature dataset to print out a nice list of your datasets,
feature classes, and attributes for each feature classes.  
"""


# See these help documents
## http://python-docx.readthedocs.org/en/latest/user/install.html
## https://python-docx.readthedocs.org/en/latest/api/style.html
## https://raw.githubusercontent.com/mikemaccana/python-docx/master/example-makedocument.py

# Use this site to use the docx module
## http://python-docx.readthedocs.org/en/latest/user/quickstart.html

# If module isn't loading, then temporarily append to sys.path variable
import os, sys

##########################################################################
# Checks that Anaconda path is in system path variable to access 64-bit Python modules
def checkPathAdded(pathToAppend):
    if pathToAppend in sys.path: 
        print 'Directory {0} already added to system path'.format(pathToAppend)
    else: 
        sys.path.append(pathToAppend)
        print 'Added {0} directory to system path'.format(pathToAppend)
        
        
#checkPathAdded(r'C:\Anaconda\Lib\site-packages')
#checkPathAdded(r'C:\Python27\ArcGISx6410.3')
#checkPathAdded(r'C:\Program Files (x86)\ArcGIS\Desktop10.3')

# Add arcpy packages from ArcGIS
#arcpy32PythonPaths = [
#'C:\\Python27\\ArcGIS10.3'
#, 'C:\\Python27\\ArcGIS10.3\\lib\\site-packages'
#, 'C:\\Program Files (x86)\\ArcGIS\\Desktop10.3\\bin'
#, 'C:\\Program Files (x86)\\ArcGIS\\Desktop10.3\\ArcPy'
#, 'C:\\Program Files (x86)\\ArcGIS\\Desktop10.3\\ArcToolBox\\Scripts'
#, 'C:\\Python27\\ArcGIS10.3\\lib\\site-packages\\win32'
#, 'C:\\Python27\\ArcGIS10.3\\lib\\site-packages\\win32\\lib']

# Paths were taken from the sys.path (PYTHONPATH variable) when running
# python 64-bit exe located C:\Python27\ArcGISx6410.3\python.exe
# If working in QGIS, these paths need to be added to sys.path
arcpy64PythonPaths = [
u'c:\\program files (x86)\\arcgis\\desktop10.3\\arcpy'
, 'C:\\windows\\system32\\python27.zip'
, 'C:\\Python27\\ArcGISx6410.3\\DLLs'
, 'C:\\Python27\\ArcGISx6410.3\\lib'
, 'C:\\Python27\\ArcGISx6410.3\\lib\\plat-win'
, 'C:\\Python27\\ArcGISx6410.3\\lib\\lib-tk'
, 'C:\\Python27\\ArcGISx6410.3'
, 'C:\\Python27\\ArcGISx6410.3\\lib\\site-packages'
, 'C:\\Program Files (x86)\\ArcGIS\\Desktop10.3\\bin64'
, 'C:\\Program Files (x86)\\ArcGIS\\Desktop10.3\\ArcPy'
, 'C:\\Program Files (x86)\\ArcGIS\\Desktop10.3\\ArcToolBox\\Scripts'
, 'C:\\Anaconda\\Lib\\site-packages']

# Adds or removes paths temporarily to sys.path variable
def addPythonPath(pathList, Add_or_Remove):
    if Add_or_Remove == 'Add':
        for eachPath in pathList:
            checkPathAdded(eachPath)
    elif Add_or_Remove == 'Remove':
        for eachPath in pathList:
            if eachPath not in sys.path: 
                continue
            else: 
                sys.path.remove(eachPath)

# Runs function to add all 64-bit python paths to the current sys.path
addPythonPath(arcpy64PythonPaths, 'Add')
###########################################################################

# Import additional modules once python path is set up 
# Import MS Word module from https://python-docx.readthedocs.org/en/latest/
from docx import Document
from docx.shared import Length
import arcpy
from datetime import datetime

##########################################
# DEFINE FUNCTIONS
##########################################

##########################################
def addWordTimestamp(wordDocObject):
    """converts the datetime.now() into strings to be accessed & added the end of logfile"""
    i = datetime.now()
    todaysDatetime = i.strftime("%m/%d/%Y %H:%M")
    # add single line in Word file for date executed
    time_para = wordDocObject.add_paragraph()
    time_para_run1 = time_para.add_run('Fields updated as of: ' + todaysDatetime)
    time_para_run1.font.italic = True
    #return time_para_run1

######################################################

def createWordDatasetHeading(wordDocObject, datasetName):
    fd_heading = wordDocObject.add_heading('{0}'.format(datasetName))
    #return fd_heading
    

######################################################

def createWordFCBulletList(wordDocObject, datasetName, featureClassName):
    # Create Feature Class Name with bulletted field names below.

    # Numbered Feature Class name, as bold
    # each new paragraph will be added after the initial one. But each new paragraph
    # needs to be referenced by a new object
    fc_para = wordDocObject.add_paragraph(style='List Number')
    ## fc_para_run1 = fc_para.add_run('referenced Feature Class' + ':')  # without reference to actual feature class
    fc_para_run1 = fc_para.add_run(featureClassName)
    fc_para_run1.font.bold = True
    fc_para_run1.font.underline = True
    
    #return fc_para_run1

######################################################

def createWordFieldNameBullets(wordDocObject, datasetName, featureClassName, field):
    """
    Create Feature Class Name with bulletted field names below. Bulleted feature class fields listed out to simplify defining
    Adding bullets under 1st paragraph with single spacing
    """
    print field.name
    
    # Assign paragraph style to Bullet List 2
    fieldName1_para = wordDocObject.add_paragraph(style = 'List Bullet 2')
    
    # Assign paragraph style to single-space, list bullet 2, bold, & underlined
    fieldName1_para.paragraph_format.line_spacing = 1.0      # make text line spacing single-spaced
    # Line Run  (of the field name)
    fieldName1_para_run1 = fieldName1_para.add_run('{0}'.format(field.name))
    fieldName1_para_run1.font.bold = True           # make text bold
    fieldName1_para_run1.font.underline = True    # make text underlined
    
    # Add in field alias name, field type, & field length in () after bold field name
    fieldName1_para_run2 = fieldName1_para.add_run(' ({0}, {1}, {2}):'.format(field.aliasName, field.type, field.length))
    fieldName1_para_run2.font.italic = True    # make text underlined
    
    #return fieldName1_para
######################################################

##########################################
# Define variables
##########################################

# Create new Word document object
document = Document()

# Establish workspace connection to SDE Database while Arcmap is open
arcpy.env.workspace = r"Database Connections\cihl-gisdat-01_sde_current_gisuser.sde"
myDir = arcpy.env.workspace

# Acess domain in workspace
domains = arcpy.da.ListDomains(arcpy.env.workspace)

# List dataset(s) in workspace
#dsList = ['sde.SDE.SteamSystem']
#dsList = arcpy.ListDatasets(feature_type='feature') # WILL ACCESS ALL FEATURE DATASETS
dsList = ['sde.SDE.Water_Distribution_Network', 'sde.SDE.Water_Distribution_Features']   # LIST OF SPECIFIC FEATURE CLASSES IN WATER UTILITIES
#dsList = ['sde.SDE.StormSewerFeatures', 'sde.SDE.StormSewerNetwork']   # LIST OF SPECIFIC FEATURE CLASSES IN STORM UTILITIES
#dsList = ['sde.SDE.Maintenance']    # DATASET FOR MAINTENANCE FEATURE CLASSES
#dsList = ['sde.SDE.SanitarySewerFeatures', 'sde.SDE.SanitarySewerNetwork']  #
#dsList = ['sde.SDE.Gas']

# Variables for saved Word file to save a new or update an existing document
metadataKeywordList = ['Gas', 'GPS', 'Maintenance', 'SantitarySewer', 'StormwaterSewer', 'Water', 'Steam']
keyword = metadataKeywordList[5]


##########################################
# Run functions
##########################################

# Print timestamp at top of page
##-->addWordTimestamp(wordDocObject)
addWordTimestamp(document)

## testing
#createWordDatasetHeading(document, dsList[0])

for ds in dsList:
    #if ds is not None: 
    print 'Dataset:', ds
    
    # run function to create a heading in Word for the GIS Dataset Name
    ##-->createWordDatasetHeading(wordDocObject, datasetNameList)
    createWordDatasetHeading(document, ds)
    
    for fc in arcpy.ListFeatureClasses(feature_dataset=ds):
        print 'Feature Class:', fc
        
        # run function to create a heading in Word for the GIS Feature Class Name in bullets
        ##->>createWordFCBulletList(wordDocObject, datasetName, featureClassName)
        createWordFCBulletList(document, ds, fc)
                
        properties = arcpy.Describe(fc)
        fields = arcpy.ListFields(fc)
        for field in fields:
            # run function to create bulleted list in Word for each field name for given feature class
            ##-->createWordFieldNameBullets(wordDocObject, datasetName, featureClassName, field)
            createWordFieldNameBullets(document, ds, fc, field)
            document.save(r"S:\GIS_Public\Documents\GISData_Workflows\GISUtilityEditingProcedures\Editing_" + keyword +r"_Fields.docx")

del document
print 'Completed export of ' + keyword + ' field lists'
            
#document.save(r"S:\GIS_Public\Documents\GISData_Workflows\GISUtilityEditingProcedures\Editing_" + keyword + r"_Fields.docx")



#for ds in dsList
#    if ds is not None: 
#        # Created new Word Document class
#        document = Document()
#
#        ## Open existing document
#        #document = Document(r'S:\GIS_Public\Documents\GISData_Workflows\GISUtilityEditingProcedures\testWordFile.docx')
#
#        # Print timestamp at top of page
#        addWordTimestamp(document)
#        print 'Dataset:', ds
#
#        # run function to create a heading in Word for the GIS Dataset Name
#        createWordDatasetHeading(document, ds)
#        
#        for fc in arcpy.ListFeatureClasses(feature_dataset=ds):
#            print 'Feature Class:', fc
#            
#            # run function to create a heading in Word for the GIS Feature Class Name in bullets
#            createWordFCBulletList(document, ds, fc)
#                    
#            properties = arcpy.Describe(fc)
#            fields = arcpy.ListFields(fc)
#                    
#            # run function to create bulleted list in Word for each field name for given feature class
#            createWordFieldNameBullets(document, ds, fc, fields)
#            
#            for field in fields:
#                createWordFieldNameBullets(document, ds, fc, field)
#                document.save(r"S:\GIS_Public\Documents\GISData_Workflows\GISUtilityEditingProcedures\Editing_Gas_Fields.docx")
        
# Saves the document for that dataset type



#            outFile.write('\n'
#                          + str(ds[i]) + ';'
#                          + str(fc) + ';'
#                          + str(properties.featureType) + ';'
#                          + str(properties.shapeType) + ';'
#                          + str(featureCount)
#                          )
#            
#            # for all fields in each fc, then write out the field attributes...
#            for field in fields:
#####################################################              
#                # INSERT CODE FOR PRINTING IN WORD #
#####################################################
#                
#                outFile.write('\n;;;;; ATTRIBUTE->;'
#                              + '{0}'.format(field.name) + ';'
#                              #+ '{0}'.format(field.aliasName) + ';' # a quick comparison of alias names showed they are generally the same
#                              + '{0}'.format(field.type) + ';'
#                              + '{0}'.format(field.length) + ';'
#                              + '{0}'.format(field.precision)+ ';'
#                              #+ '{0}'.format(field.domain)
#                              )
#
#

